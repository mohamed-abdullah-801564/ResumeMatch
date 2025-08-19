import streamlit as st
import PyPDF2
import docx
import io
from typing import Optional

def validate_file_type(file, allowed_types: list) -> bool:
    """
    Validate if the uploaded file is of allowed type
    
    Args:
        file: Streamlit uploaded file object
        allowed_types: List of allowed file extensions
    
    Returns:
        bool: True if file type is valid, False otherwise
    """
    if file is None:
        return False
    
    file_extension = file.name.split('.')[-1].lower()
    return file_extension in [ext.lower() for ext in allowed_types]

def extract_text_from_pdf(file) -> str:
    """
    Extract text from PDF file
    
    Args:
        file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    """
    try:
        # Create a file-like object from the uploaded file
        pdf_bytes = file.read()
        pdf_file = io.BytesIO(pdf_bytes)
        
        # Read PDF using PyPDF2
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        # Extract text from all pages
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")

def extract_text_from_docx(file) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    """
    try:
        # Create a file-like object from the uploaded file
        docx_bytes = file.read()
        docx_file = io.BytesIO(docx_bytes)
        
        # Read DOCX using python-docx
        doc = docx.Document(docx_file)
        text = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")

def extract_text_from_file(file) -> str:
    """
    Extract text from uploaded file based on its type
    
    Args:
        file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    """
    if file is None:
        raise Exception("No file provided")
    
    # Get file extension
    file_extension = file.name.split('.')[-1].lower()
    
    # Reset file pointer to beginning
    file.seek(0)
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    else:
        raise Exception(f"Unsupported file type: {file_extension}")

def get_file_info(file) -> dict:
    """
    Get basic information about the uploaded file
    
    Args:
        file: Streamlit uploaded file object
    
    Returns:
        dict: File information including name, size, and type
    """
    if file is None:
        return {}
    
    return {
        "name": file.name,
        "size": len(file.getvalue()),
        "type": file.type,
        "extension": file.name.split('.')[-1].lower()
    }
